import sys
import subprocess
import os
import base64
import urllib.request
import json

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches

def create_doc():
    doc = Document()
    doc.add_heading('RAG Architecture Flow Chart', 0)
    
    doc.add_paragraph(
        "Below is the target sequence diagram illustrating the True RAG flow "
        "using Azure Embeddings and ChromaDB inside your Finance AI application."
    )
    
    # Read mmd
    with open(r"d:\Finace AI\diagram.mmd", "r") as f:
        graph_def = f.read()
        
    # Standard Base64 urlsafe encode for mermaid.ink
    b64_string = base64.urlsafe_b64encode(graph_def.encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64_string}"
    
    img_path = r"d:\Finace AI\mermaid_out.png"
    print(f"Downloading from mermaid.ink...")
    try:
        urllib.request.urlretrieve(url, img_path)
    except Exception as e:
        print(f"Failed to fetch image: {e}")
        
    if os.path.exists(img_path):
        try:
            doc.add_picture(img_path, width=Inches(6.5))
            doc.add_paragraph("✅ Generated successfully.")
        except Exception as e:
            print(f"Error adding picture: {e}")
    else:
        doc.add_paragraph("Screenshot generation failed. Here is the Mermaid Diagram code instead:")
        doc.add_paragraph(graph_def)
            
    out_file = r"d:\Finace AI\RAG_Flow_Chart.docx"
    doc.save(out_file)
    print(f"\n✅ Created: {out_file}")

if __name__ == "__main__":
    create_doc()
